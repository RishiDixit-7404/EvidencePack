"""Generate the synthetic EvidencePack demo evidence folder."""

from __future__ import annotations

import csv
import shutil
from pathlib import Path

from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill
from openpyxl.utils import get_column_letter
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DEMO_DIR = ROOT / "demo_evidence"

GENERATED_FILES = [
    "access_policy_v3.pdf",
    "user_access_export_Q1.xlsx",
    "ticket_dump_march.csv",
    "control_matrix_draft.xlsx",
    "meeting_notes_kickoff.docx",
    "firewall_rules_screenshot.png",
    "controls.xlsx",
]

DOMAINS = ["Access Management", "Change Management", "Data Protection"]

CONTROLS = [
    (
        "AM-01",
        "Quarterly user access reviews are performed and documented.",
        "Policy document, Access review export, Meeting notes",
        "Access Management",
    ),
    (
        "AM-02",
        "Terminated users have access removed within 24 hours.",
        "Policy document, Ticket dump, User access export",
        "Access Management",
    ),
    (
        "AM-03",
        "Privileged or administrator access is reviewed and approved.",
        "User access export, Access review export, Meeting notes",
        "Access Management",
    ),
    (
        "AM-04",
        "Password policy requirements are formally documented.",
        "Policy document",
        "Access Management",
    ),
    (
        "CM-01",
        "Application or infrastructure changes are logged through tickets.",
        "Ticket dump",
        "Change Management",
    ),
    (
        "CM-02",
        "Changes are reviewed and approved before implementation.",
        "Ticket dump, Meeting notes",
        "Change Management",
    ),
    (
        "CM-03",
        "Emergency changes are documented after implementation.",
        "Ticket dump, Meeting notes",
        "Change Management",
    ),
    (
        "CM-04",
        "Change records include implementation status and description.",
        "Ticket dump",
        "Change Management",
    ),
    (
        "DP-01",
        "Sensitive or personal data handling requirements are documented.",
        "Policy document, Meeting notes",
        "Data Protection",
    ),
    (
        "DP-02",
        "Access to sensitive data is restricted by role or access level.",
        "User access export, Policy document",
        "Data Protection",
    ),
    (
        "DP-03",
        "Security incidents are tracked and reviewed.",
        "Ticket dump, Meeting notes",
        "Data Protection",
    ),
    (
        "DP-04",
        "Firewall or network protection evidence is maintained.",
        "Firewall rules screenshot",
        "Data Protection",
    ),
]


def _escape_pdf_text(value: str) -> str:
    return value.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def _build_pdf_stream(lines: list[str]) -> bytes:
    commands = ["BT", "/F1 11 Tf", "50 760 Td", "14 TL"]
    for line in lines:
        commands.append(f"({_escape_pdf_text(line)}) Tj")
        commands.append("T*")
    commands.append("ET")
    return ("\n".join(commands) + "\n").encode("latin-1")


def _write_text_pdf(path: Path, pages: list[list[str]]) -> None:
    objects: list[bytes] = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"",  # Pages object is filled after page object numbers are known.
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    page_object_numbers: list[int] = []

    for page_lines in pages:
        stream = _build_pdf_stream(page_lines)
        content_number = len(objects) + 1
        objects.append(
            b"<< /Length "
            + str(len(stream)).encode("ascii")
            + b" >>\nstream\n"
            + stream
            + b"endstream"
        )
        page_number = len(objects) + 1
        page_object_numbers.append(page_number)
        objects.append(
            (
                "<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
                "/Resources << /Font << /F1 3 0 R >> >> "
                f"/Contents {content_number} 0 R >>"
            ).encode("ascii")
        )

    kids = " ".join(f"{number} 0 R" for number in page_object_numbers)
    objects[1] = f"<< /Type /Pages /Kids [{kids}] /Count {len(pages)} >>".encode("ascii")

    output = bytearray(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")
    offsets = [0]
    for index, body in enumerate(objects, start=1):
        offsets.append(len(output))
        output.extend(f"{index} 0 obj\n".encode("ascii"))
        output.extend(body)
        output.extend(b"\nendobj\n")

    xref_offset = len(output)
    output.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    output.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        output.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    output.extend(
        (
            f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_offset}\n%%EOF\n"
        ).encode("ascii")
    )
    path.write_bytes(output)


def _style_header(ws) -> None:
    fill = PatternFill("solid", fgColor="D9EAF7")
    for cell in ws[1]:
        cell.font = Font(bold=True)
        cell.fill = fill
    for column_cells in ws.columns:
        max_length = max(len(str(cell.value or "")) for cell in column_cells)
        ws.column_dimensions[get_column_letter(column_cells[0].column)].width = min(
            max(max_length + 2, 14), 42
        )


def create_access_policy_pdf(path: Path) -> None:
    pages = [
        [
            "EvidencePack Demo Company",
            "User Access Policy v3",
            "Effective Date: 2026-01-01",
            "",
            "Purpose",
            "This policy defines synthetic access governance requirements for demo use.",
            "Access must be granted according to approved job responsibilities.",
            "Managers are accountable for confirming that access remains appropriate.",
            "",
            "Quarterly Access Review",
            "Application owners perform a quarterly access review for all in-scope systems.",
            "Review evidence includes reviewer name, review date, exceptions, and remediation.",
            "Privileged accounts are reviewed separately and require explicit sign-off.",
        ],
        [
            "Joiner, Mover, Leaver Requirements",
            "",
            "New access requests require an approved ticket before provisioning.",
            "Role changes require manager approval and removal of unnecessary prior access.",
            "Terminated user access must be removed within 24 hours of HR notification.",
            "The service desk records the termination access removal SLA in the ticket notes.",
            "",
            "Exceptions",
            "Any extension for legal hold or investigation must be documented by Security.",
            "Shared accounts are prohibited unless approved by the Information Security Lead.",
        ],
        [
            "Administrative Access",
            "",
            "Administrative access is limited to authorized support personnel.",
            "Privileged access assignments require business justification and owner approval.",
            "Administrators must use named accounts and may not share credentials.",
            "Access logs are retained for monitoring and incident response.",
            "",
            "Monitoring",
            "Security reviews unusual login activity and escalates suspected misuse.",
            "Incident tickets document triage, containment, and final resolution.",
        ],
        [
            "Password Policy",
            "",
            "Passwords must contain at least twelve characters.",
            "Passwords must include complexity across letters, numbers, and symbols.",
            "Default passwords must be changed before production use.",
            "Accounts lock after repeated failed authentication attempts.",
            "Password reset requests require identity verification by the service desk.",
            "",
            "Policy Maintenance",
            "This policy is reviewed annually and after material control changes.",
        ],
    ]
    _write_text_pdf(path, pages)


def create_user_access_export(path: Path) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Q1 User Access"
    ws.append(["Username", "Role", "Last Login", "Termination Date", "Access Level"])
    rows = [
        ["alex.chen", "Finance Analyst", "2026-03-29", "", "Standard"],
        ["priya.narayan", "IT Administrator", "2026-03-31", "", "Privileged"],
        ["marco.silva", "Support Engineer", "2026-03-25", "", "Standard"],
        ["jordan.lee", "Security Manager", "2026-03-30", "", "Admin"],
        ["samira.patel", "Operations Reviewer", "2026-03-22", "", "Read Only"],
        ["taylor.morgan", "Former Contractor", "2026-02-14", "2026-03-10", "Disabled"],
        ["casey.wu", "Change Coordinator", "2026-03-28", "", "Standard"],
        ["nina.ross", "Database Administrator", "2026-03-27", "", "Privileged"],
    ]
    for row in rows:
        ws.append(row)
    _style_header(ws)
    wb.save(path)


def create_ticket_dump(path: Path) -> None:
    rows = [
        {
            "Ticket ID": "IT-2026-0301",
            "Type": "Access request",
            "Date": "2026-03-02",
            "Status": "Closed",
            "Description": "Provision standard finance application access for Alex Chen.",
        },
        {
            "Ticket ID": "IT-2026-0314",
            "Type": "Termination",
            "Date": "2026-03-10",
            "Status": "Closed",
            "Description": "Disable access for Taylor Morgan within the 24-hour SLA.",
        },
        {
            "Ticket ID": "IT-2026-0320",
            "Type": "Incident",
            "Date": "2026-03-18",
            "Status": "Closed",
            "Description": "Investigate repeated failed login attempts for service portal.",
        },
        {
            "Ticket ID": "IT-2026-0328",
            "Type": "Change request",
            "Date": "2026-03-24",
            "Status": "Approved",
            "Description": "Update firewall allowlist for approved vendor integration.",
        },
        {
            "Ticket ID": "IT-2026-0331",
            "Type": "Access request",
            "Date": "2026-03-29",
            "Status": "Open",
            "Description": "Request read-only reporting access for Operations Reviewer.",
        },
    ]
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(
            handle,
            fieldnames=["Ticket ID", "Type", "Date", "Status", "Description"],
        )
        writer.writeheader()
        writer.writerows(rows)


def create_control_matrix_draft(path: Path) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Draft Matrix"
    ws.append(["Control ID", "Domain", "Control Description", "Status", "Evidence"])
    statuses = [
        ("AM-01", "Complete", "access_policy_v3.pdf; quarterly review notes pending"),
        ("AM-02", "Partial", "ticket_dump_march.csv"),
        ("AM-03", "Partial", "user_access_export_Q1.xlsx"),
        ("AM-04", "Complete", "access_policy_v3.pdf page 4"),
        ("CM-01", "Partial", "ticket_dump_march.csv"),
        ("CM-02", "Missing", ""),
        ("CM-03", "Missing", ""),
        ("CM-04", "Partial", "meeting_notes_kickoff.docx"),
        ("DP-01", "Missing", ""),
        ("DP-02", "Missing", ""),
        ("DP-03", "Partial", "firewall_rules_screenshot.png"),
        ("DP-04", "Partial", "ticket_dump_march.csv"),
    ]
    by_id = {control[0]: control for control in CONTROLS}
    for control_id, status, evidence in statuses:
        control = by_id[control_id]
        ws.append([control_id, control[3], control[1], status, evidence])
    _style_header(ws)
    wb.save(path)


def create_meeting_notes(path: Path) -> None:
    doc = Document()
    doc.add_heading("EvidencePack Demo Kickoff Meeting Notes", level=1)
    doc.add_paragraph("Date: 2026-03-31")
    doc.add_paragraph("Attendees: Demo Client Team, Internal Audit Prep Team")
    doc.add_heading("Discussion", level=2)
    items = [
        "Completed access review walkthrough for Q1 user access export and policy evidence.",
        "Reviewed open evidence items for emergency changes, restore testing, and encryption settings.",
        "Follow-up needed from the client on privileged access approvals and backup monitoring samples.",
        "Discussed data protection requirements for firewall rule reviews and sensitive data encryption.",
        "Discussed change management expectations for approval, testing, and deployment segregation.",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("Next Steps", level=2)
    doc.add_paragraph(
        "Client will provide remaining screenshots and ticket references before the pre-review checkpoint."
    )
    doc.save(path)


def create_firewall_png(path: Path) -> None:
    width, height = 900, 520
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    try:
        title_font = ImageFont.truetype("arial.ttf", 34)
        text_font = ImageFont.truetype("arial.ttf", 20)
        small_font = ImageFont.truetype("arial.ttf", 16)
    except OSError:
        title_font = ImageFont.load_default()
        text_font = ImageFont.load_default()
        small_font = ImageFont.load_default()

    draw.rectangle((0, 0, width, 64), fill=(34, 72, 112))
    draw.text((24, 16), "Firewall Rules Export", fill="white", font=title_font)
    draw.text((24, 84), "Synthetic screenshot - unsupported image evidence", fill=(70, 70, 70), font=text_font)

    headers = ["Rule", "Source", "Destination", "Port", "Status"]
    rows = [
        ["FW-1001", "10.20.0.0/16", "app.demo.local", "443", "Allowed"],
        ["FW-1002", "VPN Users", "admin.demo.local", "22", "Restricted"],
        ["FW-1003", "Vendor NAT", "api.demo.local", "8443", "Allowed"],
        ["FW-1004", "Any", "legacy.demo.local", "3389", "Disabled"],
    ]
    x_positions = [36, 180, 385, 610, 735]
    y = 140
    draw.rectangle((24, y - 12, 860, y + 32), fill=(224, 232, 240), outline=(160, 170, 180))
    for x, header in zip(x_positions, headers):
        draw.text((x, y), header, fill=(20, 20, 20), font=text_font)
    for index, row in enumerate(rows, start=1):
        y = 140 + index * 52
        fill = (248, 250, 252) if index % 2 else (238, 244, 248)
        draw.rectangle((24, y - 12, 860, y + 32), fill=fill, outline=(205, 213, 222))
        for x, value in zip(x_positions, row):
            draw.text((x, y), value, fill=(30, 30, 30), font=small_font)
    draw.rectangle((24, 426, 860, 486), outline=(205, 80, 80), width=2)
    draw.text((42, 444), "Note: image-only evidence is intentionally unextractable in v1.0.", fill=(130, 40, 40), font=text_font)
    image.save(path, "PNG")


def create_controls(path: Path) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Controls"
    ws.append(["Control ID", "Control Description", "Required Evidence Type", "Domain"])
    for control in CONTROLS:
        ws.append(list(control))
    _style_header(ws)
    wb.save(path)


def regenerate_demo_evidence() -> list[Path]:
    if DEMO_DIR.exists():
        shutil.rmtree(DEMO_DIR)
    DEMO_DIR.mkdir(parents=True)

    create_access_policy_pdf(DEMO_DIR / "access_policy_v3.pdf")
    create_user_access_export(DEMO_DIR / "user_access_export_Q1.xlsx")
    create_ticket_dump(DEMO_DIR / "ticket_dump_march.csv")
    create_control_matrix_draft(DEMO_DIR / "control_matrix_draft.xlsx")
    create_meeting_notes(DEMO_DIR / "meeting_notes_kickoff.docx")
    create_firewall_png(DEMO_DIR / "firewall_rules_screenshot.png")
    create_controls(DEMO_DIR / "controls.xlsx")

    return [DEMO_DIR / name for name in GENERATED_FILES]


def main() -> None:
    """Regenerate all synthetic demo evidence files."""
    generated = regenerate_demo_evidence()
    print("Generated demo_evidence files:")
    for path in generated:
        print(f"- {path.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
